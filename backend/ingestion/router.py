"""
Complete File Ingestion Router
Routes different file types to appropriate parsers with progress tracking
"""

from pathlib import Path
from typing import List, Dict, Any, Optional
import asyncio
import json
import hashlib
import gc
from backend.utils.mlx_utils import cleanup_model
from backend.ingestion.document_ai import DocumentAIClient


class FileIngestionRouter:
    """
    Routes files to appropriate parsers based on extension.
    Handles text, PDF, DOCX, emails, images, and videos.
    """

    SUPPORTED_TYPES = {
        ".txt": "text",
        ".md": "text",
        ".json": "text",
        ".xml": "text",
        ".csv": "text",
        ".pdf": "pdf",
        ".docx": "docx",
        ".doc": "docx",
        ".eml": "email",
        ".msg": "email",
        ".jpg": "image",
        ".jpeg": "image",
        ".png": "image",
        ".gif": "image",
        ".bmp": "image",
        ".tiff": "image",
        ".mp4": "video",
        ".mov": "video",
        ".avi": "video",
        ".mkv": "video",
        ".webm": "video",
        ".mp3": "audio",
        ".wav": "audio",
        ".m4a": "audio",
        ".flac": "audio",
    }

    def __init__(self, progress_callback: Optional[callable] = None, llm: Optional[Any] = None):
        self.progress_callback = progress_callback
        self.llm = llm
        self.doc_ai = DocumentAIClient()
        self.processed_files: List[Dict[str, Any]] = []
        self.failed_files: List[Dict[str, Any]] = []
        self.total_chunks = 0

    async def process_files(
        self, file_paths: List[str], task_id: str, target_names: List[str]
    ) -> Dict[str, Any]:
        """Process all files and extract content."""

        results = {
            "task_id": task_id,
            "target_names": target_names,
            "processed": [],
            "failed": [],
            "total_chunks": 0,
            "total_files": len(file_paths),
        }

        valid_files = [f for f in file_paths if Path(f).exists()]

        for idx, file_path in enumerate(valid_files):
            path = Path(file_path)
            file_type = self._get_file_type(path)

            self._report_progress(
                task_id, idx / len(valid_files) * 0.5, f"Processing {path.name}..."
            )

            try:
                content = await self._extract_content(path, file_type)

                if content and len(content.strip()) > 10:
                    chunks = self._chunk_text(content, path.name)

                    file_result = {
                        "file": str(path),
                        "filename": path.name,
                        "type": file_type,
                        "content": content,
                        "chunks": chunks,
                        "chunk_count": len(chunks),
                        "file_hash": self._compute_hash(path),
                    }

                    results["processed"].append(file_result)
                    results["total_chunks"] += len(chunks)

            except Exception as e:
                results["failed"].append(
                    {"file": str(path), "filename": path.name, "error": str(e)}
                )

        self.processed_files = results["processed"]
        self.failed_files = results["failed"]
        self.total_chunks = results["total_chunks"]

        return results

    def _get_file_type(self, path: Path) -> str:
        """Determine file type from extension."""
        ext = path.suffix.lower()
        return self.SUPPORTED_TYPES.get(ext, "unknown")

    async def _extract_content(self, path: Path, file_type: str) -> str:
        """Extract text content based on file type."""

        if file_type == "text":
            return await self._extract_text(path)
        elif file_type == "pdf":
            return await self._extract_pdf(path)
        elif file_type == "docx":
            return await self._extract_docx(path)
        elif file_type == "email":
            return await self._extract_email(path)
        elif file_type == "image":
            return await self._extract_image(path)
        elif file_type in ("video", "audio"):
            return await self._extract_media(path, file_type)
        else:
            return f"[Unsupported file type: {path.suffix}]"

    async def _extract_text(self, path: Path) -> str:
        """Extract plain text."""
        encodings = ["utf-8", "latin-1", "cp1252", "iso-8859-1"]

        for encoding in encodings:
            try:
                return path.read_text(encoding=encoding)
            except UnicodeDecodeError:
                continue

        with open(path, "rb") as f:
            raw = f.read()
            return raw.decode("utf-8", errors="replace")

    async def _extract_pdf(self, path: Path) -> str:
        """Extract text from PDF (prefers Document AI if enabled)."""
        if self.doc_ai.is_enabled():
            result = await self.doc_ai.process_document(str(path), mime_type="application/pdf")
            if result:
                return f"[Document AI Extraction]\n{result['text']}"

        try:
            from pypdf import PdfReader

            reader = PdfReader(str(path))
            text_parts = []

            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text:
                    text_parts.append(f"[Page {i + 1}]\n{text}")

            return "\n\n".join(text_parts) if text_parts else "[No text found in PDF]"
        except ImportError:
            return "[PDF processing requires pypdf: pip install pypdf]"
        except Exception as e:
            return f"[PDF extraction failed: {e}]"

    async def _extract_docx(self, path: Path) -> str:
        """Extract text from DOCX."""
        try:
            from docx import Document

            doc = Document(str(path))
            paragraphs = []

            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)

            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    if any(cells):
                        tables_text.append(" | ".join(cells))

            result = "\n".join(paragraphs)
            if tables_text:
                result += "\n\n[Tables]\n" + "\n".join(tables_text)

            return result or "[Empty document]"
        except ImportError:
            return "[DOCX processing requires python-docx: pip install python-docx]"
        except Exception as e:
            return f"[DOCX extraction failed: {e}]"

    async def _extract_email(self, path: Path) -> str:
        """Extract content from email files."""
        try:
            from email import policy
            from email.parser import BytesParser

            with open(path, "rb") as f:
                msg = BytesParser(policy=policy.default).parse(f)

            parts = []
            parts.append(f"Subject: {msg.get('Subject', 'No Subject')}")
            parts.append(f"From: {msg.get('From', 'Unknown')}")
            parts.append(f"To: {msg.get('To', 'Unknown')}")
            parts.append(f"Date: {msg.get('Date', 'Unknown')}")
            parts.append("")

            body = msg.get_body(preferencelist=("plain", "html"))
            if body:
                content = body.get_content()
                parts.append(f"\n--- Body ---\n{content}")

            return "\n".join(parts)
        except Exception as e:
            return f"[Email extraction failed: {e}]"

    async def _extract_image(self, path: Path) -> str:
        """Extract text from images (prefers Document AI if enabled)."""
        if self.doc_ai.is_enabled():
            # Basic map for common image types
            mime_map = {".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg", ".tiff": "image/tiff"}
            mime_type = mime_map.get(path.suffix.lower(), "image/jpeg")

            result = await self.doc_ai.process_document(str(path), mime_type=mime_type)
            if result:
                return f"[Document AI OCR]\n{result['text']}"

        try:
            import pytesseract
            from PIL import Image

            img = Image.open(str(path))
            text = pytesseract.image_to_string(img)

            return text.strip() if text.strip() else "[No text detected in image]"
        except ImportError:
            return "[Image OCR requires pytesseract and Pillow]"
        except Exception as e:
            return f"[Image OCR failed: {e}]"

    async def _extract_media(self, path: Path, media_type: str) -> str:
        """Extract content from media (uses Gemini for video if available)."""
        if media_type == "video" and self.llm:
            try:
                # Use Gemini 1.5 Flash for video analysis
                with open(path, "rb") as f:
                    video_data = f.read()

                response = await self.llm.generate(
                    contents=[
                        {"mime_type": "video/mp4", "data": video_data},
                        "Describe this video in detail for a legal case. Identify key events, timestamps, and entities."
                    ],
                    task_complexity="analysis"
                )
                # Handle potential wrapper result
                text_content = getattr(response, 'text', getattr(response, 'content', str(response)))
                return f"[Gemini Video Analysis]\n{text_content}"
            except Exception as e:
                print(f"Gemini video analysis failed: {e}")

        try:
            if media_type == "audio":
                return await self._transcribe_audio(path)
            else:
                audio_path = await self._extract_audio_from_video(path)
                if audio_path:
                    result = await self._transcribe_audio(audio_path)
                    try:
                        Path(audio_path).unlink()
                    except:
                        pass
                    return result
                return "[Failed to extract audio from video]"
        except Exception as e:
            return f"[Media transcription failed: {e}]"

    async def _extract_audio_from_video(self, video_path: Path) -> Optional[str]:
        """Extract audio track from video file."""
        try:
            from moviepy.editor import VideoFileClip

            temp_audio = video_path.with_suffix(".wav")

            video = VideoFileClip(str(video_path))
            if video.audio:
                video.audio.write_audiofile(str(temp_audio), verbose=False, logger=None)
                video.close()
                return str(temp_audio)
            else:
                video.close()
                return None
        except ImportError:
            return None
        except Exception:
            return None

    async def _transcribe_audio(self, audio_path: Path) -> str:
        """Transcribe audio using MLX-Whisper or Gemini Flash fallback."""
        from backend.config.settings import settings

        if settings.DISABLE_LOCAL_MLX and self.llm:
            try:
                # Use Gemini 1.5 Flash for audio transcription
                with open(audio_path, "rb") as f:
                    audio_data = f.read()

                response = await self.llm.generate(
                    contents=[
                        {"mime_type": "audio/wav", "data": audio_data},
                        {"text": "Transcribe this audio file accurately. Return only the transcription text."}
                    ],
                    task_complexity="analysis"
                )
                text_content = getattr(response, 'text', getattr(response, 'content', str(response)))
                return f"[Gemini Flash Transcription]\n{text_content}"
            except Exception as e:
                print(f"Gemini audio transcription failed: {e}")
                # Fallback to local if flash fails and privacy is not strict
                if settings.DATA_PRIVACY_STRICT:
                    return f"[Transcription failed: {e}]"

        try:
            import whisper
            from backend.utils.mlx_utils import cleanup_model

            model = whisper.load_model("base")
            result = model.transcribe(str(audio_path))

            text = result["text"].strip() or "[No speech detected]"

            # Use centralized memory cleanup
            cleanup_model(model)

            return text
        except ImportError:
            return "[Audio transcription requires whisper: pip install whisper]"
        except Exception as e:
            return f"[Transcription failed: {e}]"

    def _chunk_text(
        self, text: str, source: str, chunk_size: int = 1000
    ) -> List[Dict[str, Any]]:
        """Split text into overlapping chunks for vector storage."""

        chunks = []
        start = 0
        text_length = len(text)

        while start < text_length:
            end = min(start + chunk_size, text_length)
            chunk_text = text[start:end]

            chunks.append(
                {
                    "text": chunk_text,
                    "source": source,
                    "start_index": start,
                    "end_index": end,
                    "chunk_id": f"{source}_{start}",
                }
            )

            start = end
            if start >= text_length:
                break

        return chunks

    def _compute_hash(self, path: Path) -> str:
        """Compute SHA256 hash of file."""
        sha256 = hashlib.sha256()

        with open(path, "rb") as f:
            for chunk in iter(lambda: f.read(8192), b""):
                sha256.update(chunk)

        return sha256.hexdigest()[:16]

    def _report_progress(self, task_id: str, progress: float, message: str):
        """Report progress to callback if available."""
        if self.progress_callback:
            self.progress_callback(task_id, progress, message)


async def process_files(
    file_paths: List[str], task_id: str, target_names: List[str]
) -> Dict[str, Any]:
    """Standalone function for file processing."""
    router = FileIngestionRouter()
    return await router.process_files(file_paths, task_id, target_names)
